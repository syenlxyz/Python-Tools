from alive_progress import alive_it
from datetime import datetime
from docx import Document
from pathlib import Path
import pandas as pd
import re
import shutil

Less_Than_Twenty = {
    0: '',
    1: 'One',
    2: 'Two',
    3: 'Three',
    4: 'Four',
    5: 'Five',
    6: 'Six',
    7: 'Seven',
    8: 'Eight',
    9: 'Nine',
    10: 'Ten',
    11: 'Eleven',
    12: 'Twelve',
    13: 'Thirteen',
    14: 'Fourteen',
    15: 'Fifteen',
    16: 'Sixteen',
    17: 'Seventeen',
    18: 'Eighteen',
    19: 'Nineteen'
}

Less_Than_Hundred = {
    2: 'Twenty',
    3: 'Thirty',
    4: 'Forty',
    5: 'Fifty',
    6: 'Sixty',
    7: 'Seventy',
    8: 'Eighty',
    9: 'Ninety'
}

More_Than_Thousand = {
    1: 'Thousand',
    2: 'Million',
    3: 'Biliion',
    4: 'Trillion'
}

def run():
    input_path = Path.cwd() / 'input'
    if not input_path.is_dir():
        input_path.mkdir()
    
    output_path = Path.cwd() / 'output'
    if not output_path.is_dir():
        output_path.mkdir()
    else:
        shutil.rmtree(str(output_path))
        output_path.mkdir()
    
    template_path = Path.cwd() / 'template.docx'
    if not template_path.exists():
        print(f'{template_path.name} does not exist')
        return None
    
    options = {
        'length': 70,
        'spinner': 'classic',
        'bar': 'classic2',
        'receipt_text': True,
        'dual_line': True
    }
    
    file_list = list(input_path.glob('**/*.xlsx'))
    results = alive_it(
        file_list, 
        len(file_list), 
        finalize=lambda bar: bar.text(f': done'),
        **options
    )

    for file_path in results:
        results.text(f':{file_path.name}')
        table = get_table(file_path)
        for index, data in enumerate(table):
            target_path = output_path / f'{file_path.stem}-{index + 1}.docx'
            create_voucher(template_path, target_path, data)

def create_voucher(template_path, target_path, data):
    doc = Document(template_path)
    table = doc.tables[0]
    for row in table.rows:
        for item in row.cells:
            pattern = r'\$\{(.*?)\}'
            result = re.findall(pattern, item.text)
            if result:
                key = result[0]
                value = data[key]
                for index, run in enumerate(item.paragraphs[0].runs):
                    if index == 0:
                        run.text = str(value)
                    else:
                        run.text = ''
    doc.save(target_path)

def get_table(file_path):
    df = pd.read_excel(file_path)
    
    amounts = ['Amount' + str(i + 1) for i in range(8)]
    for index, amount in enumerate(amounts):
        if index == 0:
            df['Total'] = df[amount]
        elif not df[amount].isna().sum():
            df['Total'] = df['Total'] + df[amount]
    df['Ringgit'] = df['Total'].apply(num_to_word)
    
    df['Date'] = df['Date'].dt.strftime('%d/%m/%Y')
    for amount in amounts:
        if not df[amount].isna().sum():
            df[amount] = df[amount].apply(lambda row: f'{row:,.2f}')
    df['Total'] = df['Total'].apply(lambda row: f'{row:,.2f}')
    df = df.fillna('')
    
    table = df.to_dict('records')
    return table

def num_to_word(num):
    if type(num) == float:
        decimal = round(num - int(num), 2)
        if decimal:
            text = ' '.join([num_to_word(int(num)), 'and Cents', num_to_word(int(decimal * 100)), 'Only'])
        else:
            text = ' '.join([num_to_word(int(num)), 'Only'])
        text = text.replace('  ', ' ')
        return text
    num_digit = len(str(num))
    power = (num_digit - 1) // 3
    if num < 20:
        text = Less_Than_Twenty[num]
        return text
    elif num < 100:
        text = ' '.join([Less_Than_Hundred[num // 10], num_to_word(num % 10)])
        return text
    elif num < 1000:
        text = ' '.join([num_to_word(num // 100), 'Hundred', num_to_word(num % 100)])
        return text
    elif power < 5:
        text = ' '.join([num_to_word(num // 1000 ** power), More_Than_Thousand[power], num_to_word(num % 1000 ** power)])
        return text

if __name__ == '__main__':
    print(f'Running {Path(__file__).parent.name}')
    start_time = datetime.now()
    run()
    end_time = datetime.now()
    run_time = end_time - start_time
    print(f'Execution time: {run_time}')