from alive_progress import alive_it
from datetime import datetime
from docx import Document
from pathlib import Path
import pandas as pd
import re
import shutil

Less_Than_Twenty = {
    0: '',
    1: 'One',
    2: 'Two',
    3: 'Three',
    4: 'Four',
    5: 'Five',
    6: 'Six',
    7: 'Seven',
    8: 'Eight',
    9: 'Nine',
    10: 'Ten',
    11: 'Eleven',
    12: 'Twelve',
    13: 'Thirteen',
    14: 'Fourteen',
    15: 'Fifteen',
    16: 'Sixteen',
    17: 'Seventeen',
    18: 'Eighteen',
    19: 'Nineteen'
}

Less_Than_Hundred = {
    2: 'Twenty',
    3: 'Thirty',
    4: 'Forty',
    5: 'Fifty',
    6: 'Sixty',
    7: 'Seventy',
    8: 'Eighty',
    9: 'Ninety'
}

More_Than_Thousand = {
    1: 'Thousand',
    2: 'Million',
    3: 'Biliion',
    4: 'Trillion'
}

def run():
    input_path = Path.cwd() / 'input'
    if not input_path.is_dir():
        input_path.mkdir()
    
    output_path = Path.cwd() / 'output'
    if not output_path.is_dir():
        output_path.mkdir()
    else:
        shutil.rmtree(str(output_path))
        output_path.mkdir()
    
    journal_voucher = Path.cwd() / 'journal-voucher.docx'
    if not journal_voucher.exists():
        print(f'{journal_voucher.name} does not exist')
    
    payment_voucher = Path.cwd() / 'payment-voucher.docx'
    if not payment_voucher.exists():
        print(f'{payment_voucher.name} does not exist')
    
    receipt_voucher = Path.cwd() / 'receipt-voucher.docx'
    if not receipt_voucher.exists():
        print(f'{receipt_voucher.name} does not exist')
    
    options = {
        'length': 70,
        'spinner': 'classic',
        'bar': 'classic2',
        'receipt_text': True,
        'dual_line': True
    }
    
    file_list = list(input_path.glob('**/*.xlsx'))
    results = alive_it(
        file_list, 
        len(file_list), 
        finalize=lambda bar: bar.text(f'Generating Voucher: done'),
        **options
    )

    for file_path in results:
        results.text(f'Generating Voucher: {file_path.name}')
        table = get_table(file_path)
        for index, data in enumerate(table):
            if data['Type'] == 'JV':
                template = journal_voucher
            if data['Type'] == 'PV':
                template = payment_voucher
            if data['Type'] == 'RV':
                template = receipt_voucher
            target_path = output_path / f'{file_path.stem}-{index + 1}.docx'
            create_voucher(template, data, target_path)

def get_table(file_path):
    df = pd.read_excel(file_path)
    df = df.fillna('')
    
    df['Date'] = df.apply(lambda row: get_date_string(row), axis='columns')
    df['Voucher'] = df.apply(lambda row: get_voucher_string(row), axis='columns')
    
    for index in range(12):
        df['Account' + str(index + 1)] = df.apply(lambda row: get_account_string(row, index + 1), axis='columns')
        df['Particular' + str(index + 1)] = df.apply(lambda row: get_particular_string(row, index + 1), axis='columns')
        df['AmountTotal'] = df.apply(lambda row: get_amount_total(row, 'Amount', index + 1), axis='columns')
        df['DebitTotal'] = df.apply(lambda row: get_amount_total(row, 'Debit', index + 1), axis='columns')
        df['CreditTotal'] = df.apply(lambda row: get_amount_total(row, 'Credit', index + 1), axis='columns')
        df['Amount' + str(index + 1)] = df.apply(lambda row: get_amount_string(row, 'Amount', index + 1), axis='columns')
        df['Debit' + str(index + 1)] = df.apply(lambda row: get_amount_string(row, 'Debit', index + 1), axis='columns')
        df['Credit' + str(index + 1)] = df.apply(lambda row: get_amount_string(row, 'Credit', index + 1), axis='columns')
    
    
    df['Ringgit'] = df.apply(lambda row: get_ringgit_string(row), axis='columns')
    df['AccountTotal'] = df.apply(lambda row: get_account_string(row, 'Total'), axis='columns')
    df['AmountTotal'] = df.apply(lambda row: get_amount_string(row, 'Amount', 'Total'), axis='columns')
    df['DebitTotal'] = df.apply(lambda row: get_amount_string(row, 'Debit', 'Total'), axis='columns')
    df['CreditTotal'] = df.apply(lambda row: get_amount_string(row, 'Credit', 'Total'), axis='columns')
    
    table = df.to_dict('records')
    return table

def get_date_string(row):
    day = str(row['Day']).zfill(2)
    month = str(row['Month']).zfill(2)
    year = str(row['Year'])
    date_string = '/'.join([day, month, year])
    return date_string

def get_voucher_string(row):
    bank = row['Bank']
    voucher_type = row['Type']
    sequence = str(row['Sequence']).zfill(2)
    month = str(row['Month']).zfill(2)
    voucher_string = f'{bank} {voucher_type} {sequence}/{month}'
    return voucher_string

def get_account_string(row, index):
    account = row['Account' + str(index)]
    voucher_type = row['Type']
    if account:
        if voucher_type == 'PV':
            if index == 'Total':
                account_string = f'CR {account}'
            else:
                account_string = f'DR {account}'
        if voucher_type == 'RV':
            if index == 'Total':
                account_string = f'DR {account}'
            else:
                account_string = f'CR {account}'
        if voucher_type == 'JV':
            debit = row['Debit' + str(index)]
            credit = row['Crebit' + str(index)]
            if debit:
                account_string = f'DR {account}'
            if credit:
                account_string = f'CR {account}'
    else:
        account_string = ''
    return account_string

def get_particular_string(row, index):
    description = row['Description' + str(index)]
    period = row['Period' + str(index)]
    if description and period:
        particular_string = f'{description} ({period})'
    elif description:
        particular_string = description
    else:
        particular_string = ''
    return particular_string

def get_amount_total(row, name, index):
    column = row[name + str(index)]
    total = row[name + 'Total']
    if column:
        if index == 1:
            amount_total = column
        else:
            amount_total = total + column
    else:
        amount_total = total
    return amount_total

def get_amount_string(row, name, index):
    column = row[name + str(index)]
    if column:
        amount_string = f'RM{column:,.2f}'
    else:
        amount_string = ''
    return amount_string

def get_ringgit_string(row):
    total = row['AmountTotal']
    if total:
        total = float(total)
        ringgit_string = num_to_word(total)
        return ringgit_string

def num_to_word(num):
    if type(num) == float:
        decimal = round(num - int(num), 2)
        if decimal:
            text = ' '.join([num_to_word(int(num)), 'and Cents', num_to_word(int(decimal * 100)), 'Only'])
        else:
            text = ' '.join([num_to_word(int(num)), 'Only'])
        text = text.replace('  ', ' ')
        return text
    num_digit = len(str(num))
    power = (num_digit - 1) // 3
    if num < 20:
        text = Less_Than_Twenty[num]
        return text
    elif num < 100:
        text = ' '.join([Less_Than_Hundred[num // 10], num_to_word(num % 10)])
        return text
    elif num < 1000:
        text = ' '.join([num_to_word(num // 100), 'Hundred', num_to_word(num % 100)])
        return text
    elif power < 5:
        text = ' '.join([num_to_word(num // 1000 ** power), More_Than_Thousand[power], num_to_word(num % 1000 ** power)])
        return text

def create_voucher(template, data, target_path):
    doc = Document(template)
    table = doc.tables[0]
    for row in table.rows:
        for item in row.cells:
            pattern = r'\$\{(.*?)\}'
            result = re.findall(pattern, item.text)
            if result:
                key = result[0]
                value = data[key]
                for index, run in enumerate(item.paragraphs[0].runs):
                    if index == 0:
                        run.text = str(value)
                    else:
                        run.text = ''
    doc.save(target_path)

if __name__ == '__main__':
    print(f'Running {Path(__file__).parent.name}')
    start_time = datetime.now()
    run()
    end_time = datetime.now()
    run_time = end_time - start_time
    print(f'Execution time: {run_time}')